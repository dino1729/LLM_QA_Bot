"""
Simple vector store with JSON persistence and cosine similarity search.
Replaces LlamaIndex VectorStoreIndex for lightweight RAG operations.
"""
import json
import logging
import os
import re
import tempfile
import uuid
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

import numpy as np

logger = logging.getLogger(__name__)


class Document:
    """Simple document with text and metadata."""

    def __init__(self, text: str, metadata: Dict[str, Any] = None, doc_id: str = None):
        self.text = text
        self.metadata = metadata or {}
        self.doc_id = doc_id or str(uuid.uuid4())

    def get_content(self) -> str:
        return self.text


class RetrievalResult:
    """Result from a similarity search."""

    def __init__(self, text: str, metadata: Dict[str, Any], score: float, doc_id: str):
        self.text = text
        self.metadata = metadata
        self.score = score
        self.doc_id = doc_id

    def get_content(self) -> str:
        return self.text


def cosine_similarity(a: List[float], b: List[float]) -> float:
    """Compute cosine similarity between two vectors."""
    a_arr = np.array(a)
    b_arr = np.array(b)
    norm_a = np.linalg.norm(a_arr)
    norm_b = np.linalg.norm(b_arr)
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return float(np.dot(a_arr, b_arr) / (norm_a * norm_b))


def chunk_text(text: str, chunk_size: int = 1024, overlap: int = 200) -> List[str]:
    """
    Split text into overlapping chunks at sentence boundaries.

    Args:
        text: Text to split
        chunk_size: Target chunk size in characters (must be > 0)
        overlap: Number of overlap characters between chunks (clamped to < chunk_size)
    """
    if not text.strip():
        return []
    chunk_size = max(chunk_size, 1)
    overlap = max(0, min(overlap, chunk_size - 1))
    sentences = re.split(r'(?<=[.!?])\s+', text)
    if not sentences:
        return []

    chunks = []
    current_chunk = []
    current_size = 0

    for sentence in sentences:
        sentence_len = len(sentence)

        if current_size + sentence_len > chunk_size and current_chunk:
            chunks.append(' '.join(current_chunk))
            # Keep sentences for overlap
            overlap_text = ' '.join(current_chunk)
            while len(overlap_text) > overlap and current_chunk:
                current_chunk.pop(0)
                overlap_text = ' '.join(current_chunk)
            current_size = len(overlap_text)

        current_chunk.append(sentence)
        current_size += sentence_len + (1 if len(current_chunk) > 1 else 0)

    if current_chunk:
        chunks.append(' '.join(current_chunk))

    return chunks if chunks else [text]


def read_documents_from_directory(directory: str) -> List[Document]:
    """
    Read all supported files from a directory and return Documents.
    Handles .txt, .md, .pdf, and .docx files.
    """
    documents = []
    dir_path = Path(directory)

    if not dir_path.exists():
        return documents

    for file_path in sorted(dir_path.iterdir()):
        if not file_path.is_file():
            continue

        ext = file_path.suffix.lower()
        text = None

        if ext in ('.txt', '.md', '.csv'):
            try:
                text = file_path.read_text(errors='ignore')
            except Exception as e:
                logger.warning(f"Failed to read {file_path}: {e}")

        elif ext == '.pdf':
            try:
                import pypdf
                reader = pypdf.PdfReader(str(file_path))
                text = '\n'.join(page.extract_text() or '' for page in reader.pages)
            except ImportError:
                logger.warning(f"Cannot read PDF {file_path}: pypdf not installed")
            except Exception as e:
                logger.warning(f"Failed to read PDF {file_path}: {e}")

        elif ext == '.docx':
            try:
                import docx
                doc = docx.Document(str(file_path))
                text = '\n'.join(p.text for p in doc.paragraphs)
            except ImportError:
                logger.warning(f"Cannot read DOCX {file_path}: python-docx not installed")
            except Exception as e:
                logger.warning(f"Failed to read DOCX {file_path}: {e}")

        else:
            # Try reading as text for unknown extensions
            try:
                text = file_path.read_text(errors='ignore')
            except Exception as e:
                logger.warning(f"Failed to read {file_path} with unknown extension '{ext}': {e}")
                continue

        if text and text.strip():
            documents.append(Document(
                text=text,
                metadata={"filename": file_path.name, "filepath": str(file_path)}
            ))

    return documents


class SimpleVectorStore:
    """
    JSON-backed vector store with cosine similarity search.
    Drop-in replacement for LlamaIndex VectorStoreIndex.
    """

    def __init__(self, persist_dir: str, embed_fn: Callable = None):
        """
        Args:
            persist_dir: Directory for JSON persistence
            embed_fn: Function (text, input_type="passage") -> List[float]
        """
        self.persist_dir = Path(persist_dir)
        self.persist_dir.mkdir(parents=True, exist_ok=True)
        self.embed_fn = embed_fn
        self._documents: Dict[str, Dict] = {}
        self._load()

    def _store_file(self) -> Path:
        return self.persist_dir / "vector_store.json"

    def _load(self):
        """Load from disk."""
        store_file = self._store_file()
        if store_file.exists():
            try:
                with open(store_file, encoding="utf-8") as f:
                    self._documents = json.load(f)
            except (json.JSONDecodeError, IOError) as e:
                logger.warning(f"Failed to load vector store from {store_file}: {e}")
                self._documents = {}

    def persist(self):
        """Save to disk atomically via temp file + os.replace()."""
        store_file = self._store_file()
        fd, tmp_path = tempfile.mkstemp(dir=str(self.persist_dir), suffix=".tmp")
        try:
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                json.dump(self._documents, f)
            os.replace(tmp_path, str(store_file))
        except Exception:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
            raise

    def insert(self, doc: Document, persist: bool = True):
        """Insert a single document with its embedding.

        Note: This method persists to disk on every call. For batch
        operations, use ``insert_many()`` to avoid repeated writes.
        """
        embedding = None
        if self.embed_fn:
            embedding = self.embed_fn(doc.text, input_type="passage")

        self._documents[doc.doc_id] = {
            "text": doc.text,
            "metadata": doc.metadata,
            "embedding": embedding
        }
        if persist:
            self.persist()

    def insert_many(self, docs: List[Document]):
        """Insert multiple documents."""
        for doc in docs:
            embedding = None
            if self.embed_fn:
                embedding = self.embed_fn(doc.text, input_type="passage")
            self._documents[doc.doc_id] = {
                "text": doc.text,
                "metadata": doc.metadata,
                "embedding": embedding
            }
        self.persist()

    def search(self, query: str, top_k: int = 5) -> List[RetrievalResult]:
        """Search by cosine similarity (O(n) linear scan over all documents)."""
        if not self._documents or not self.embed_fn:
            return []

        query_embedding = self.embed_fn(query, input_type="query")

        scored = []
        for doc_id, data in self._documents.items():
            if data.get("embedding"):
                score = cosine_similarity(query_embedding, data["embedding"])
                scored.append((doc_id, data, score))

        scored.sort(key=lambda x: x[2], reverse=True)

        return [
            RetrievalResult(
                text=data["text"],
                metadata=data["metadata"],
                score=score,
                doc_id=doc_id
            )
            for doc_id, data, score in scored[:top_k]
        ]

    def get_document(self, doc_id: str) -> Optional[Document]:
        """Get a document by ID."""
        data = self._documents.get(doc_id)
        if data:
            return Document(text=data["text"], metadata=data["metadata"], doc_id=doc_id)
        return None

    def delete_document(self, doc_id: str) -> bool:
        """Delete a document by ID."""
        if doc_id in self._documents:
            del self._documents[doc_id]
            self.persist()
            return True
        return False

    @property
    def docs(self) -> Dict[str, Any]:
        """Return all document data keyed by ID."""
        return dict(self._documents)

    def __len__(self):
        return len(self._documents)

    @classmethod
    def from_documents(cls, documents: List[Document], persist_dir: str,
                       embed_fn: Callable = None) -> 'SimpleVectorStore':
        """Create a new store from a list of documents."""
        store = cls(persist_dir=persist_dir, embed_fn=embed_fn)
        if documents:
            store.insert_many(documents)
        return store


class StreamingResult:
    """Wrapper to provide .response_gen for backwards compatibility with FastAPI endpoints."""

    def __init__(self, generator):
        self.response_gen = generator
